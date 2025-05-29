from abc import (
    ABC,
    abstractmethod,
)
from pathlib import Path
from typing import Literal
import logging

from pypdf import (
    PdfReader,
    DocumentInformation,
)
from bs4 import BeautifulSoup
import langdetect
import docx
import openpyxl


class ContentHandler(ABC):
    """
    Базовый класс обработчика, который имеет основной метод handle и абстрактный _handle.
    Основной метод обрабатывает общие исключения и возможные ситуации, и также вызывает метод _handle,
    который обязательно должен быть реализован в наследниках, где способ обработки может быть любым.

    Однако не учтен момент, что между наследниками и базовым классом нет соглашения о поле self.metadata.
    """
    def __init__(self, dest_folder: str = "."):
        self.dest_folder: str = dest_folder
        self.download_status: Literal["success", "failed_processing"] = "success"
        self.error_message: str = ""
        self.path: str = ""
        self.metadata: dict = {}

    def handle(
        self,
        file_path: str,
        *,
        dest_folder: str | None = None,
        file_name: str | None = None,
    ):
        dest_folder: Path = self._mkdir(dest_folder or self.dest_folder)
        file_name: str = file_name or file_path.split("/")[-1].split(".")[0] + ".txt"
        dest_file_path = str(dest_folder / file_name)

        try:
            logging.info(f"Starting new processing {file_path}")
            self._handle(file_path, dest_file_path)
        except Exception as e:
            logging.warning(f"Processing failed: {e}")
            self.download_status = "failed_processing"
            self.error_message = str(e)

        self.path = dest_file_path
        return dest_file_path

    @abstractmethod
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ): ...

    def _mkdir(self, folder: str) -> Path:
        folder: Path = Path(folder)
        folder.mkdir(parents=True, exist_ok=True)
        return folder


class PDFHandler(ContentHandler):
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ):
        pdf_reader = PdfReader(file_path)
        metadata: DocumentInformation = pdf_reader.metadata
        text: str = "\n".join(page.extract_text() for page in pdf_reader.pages)

        creation_date = metadata.creation_date
        self.metadata = {
            "document_page_count": len(pdf_reader.pages),
            "author": metadata.author,
            "creation_date": creation_date.strftime("%Y-%m-%d %H:%M:%S") if creation_date else None,
            "language": langdetect.detect(text),
        }

        with open(dest_file_path, "w") as file:
            file.write(text)


class DocXHandler(ContentHandler):
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ):
        document = docx.Document(file_path)
        metadata = document.core_properties
        text: str = "\n".join(paragraph.text for paragraph in document.paragraphs)

        creation_date = metadata.created
        self.metadata = {
            "document_page_count": len(document.paragraphs),
            "author": metadata.author,
            "creation_date": creation_date.strftime("%Y-%m-%d %H:%M:%S") if creation_date else None,
            "language": langdetect.detect(text),
        }

        with open(dest_file_path, "w") as file:
            file.write(text)


class XLSXHandler(ContentHandler):
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ):
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        metadata = workbook.properties
        text: str = "\n".join(
            " ".join(str(cell) for cell in row if cell is not None)
            for sheet in workbook.worksheets
            for row in sheet.iter_rows(values_only=True)
        )

        creation_date = metadata.created
        self.metadata = {
            "author": metadata.creator,
            "creation_date": creation_date.strftime("%Y-%m-%d %H:%M:%S") if creation_date else None,
            "language": langdetect.detect(text),
        }

        with open(dest_file_path, "w") as file:
            file.write(text)


class PageHandler(ContentHandler):
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ):
        with open(file_path, "r") as input_file:
            soup = BeautifulSoup(input_file, "html.parser")

            self.metadata = {"language": langdetect.detect(soup.text)}

            with open(dest_file_path, "w") as output_file:
                output_file.write(soup.text)
